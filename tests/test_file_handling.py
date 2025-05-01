import pytest
from pathlib import Path
import magic
import io
from PIL import Image

from src.utils.file_validator import FileValidator, InvalidFileTypeError
from src.utils.exceptions import FileCorruptError, FileAccessError, FileTooLargeError

class TestFileHandling:
    """
    Test suite for file handling functionality.
    Tests file type detection, validation, and basic content checking.
    """

    def test_valid_pdf_detection(self, tmp_path):
        """
        Test that a valid PDF file is correctly identified.
        Uses a temporary file to ensure test isolation.
        """
        # Create a simple PDF file for testing
        test_file = tmp_path / "test.pdf"
        test_file.write_bytes(b"%PDF-1.5\n%\x93\x8C\x8B\x9E")  # Basic PDF header

        validator = FileValidator()
        assert validator.is_valid_file_type(test_file) == True
        assert validator.get_file_type(test_file) == "application/pdf"

    def test_valid_image_detection(self, tmp_path):
        """
        Test that valid image files are correctly identified.
        """
        # Create a small valid PNG image
        img = Image.new('RGB', (100, 100), color='red')
        png_path = tmp_path / "test.png"
        img.save(png_path)

        # Create a small valid JPEG image
        jpg_path = tmp_path / "test.jpg"
        img.save(jpg_path)

        validator = FileValidator()
        # Test PNG
        assert validator.is_valid_file_type(png_path) == True
        assert validator.get_file_type(png_path) == "image/png"

        # Test JPEG
        assert validator.is_valid_file_type(jpg_path) == True
        assert validator.get_file_type(jpg_path) == "image/jpeg"

    def test_invalid_file_type(self, tmp_path):
        """
        Test that files with unsupported types are rejected.
        """
        # Create a text file with .pdf extension
        test_file = tmp_path / "test.xyz"
        test_file.write_text("This is not a valid file type")

        validator = FileValidator()
        with pytest.raises(InvalidFileTypeError, match="Unsupported file type"):
            validator.is_valid_file_type(test_file)

    def test_empty_file(self, tmp_path):
        """
        Test handling of empty files.
        Empty files should be rejected.
        """
        test_file = tmp_path / "empty.pdf"
        test_file.touch()  # Creates an empty file

        validator = FileValidator()
        with pytest.raises(InvalidFileTypeError, match="Empty file"):
            validator.is_valid_file_type(test_file)

    def test_file_size_limit(self, tmp_path):
        """
        Test that files exceeding size limit are rejected.
        """
        test_file = tmp_path / "large.pdf"
        # Create a file that exceeds our limit (let's say 10MB)
        test_file.write_bytes(b"%PDF-1.5" + b"0" * (11 * 1024 * 1024))  # 11MB file

        validator = FileValidator()
        with pytest.raises(FileTooLargeError, match="File exceeds maximum allowed size"):
            validator.is_valid_file_type(test_file)

    def test_nonexistent_file(self):
        """
        Test handling of nonexistent files.
        """
        validator = FileValidator()
        with pytest.raises(FileAccessError, match="File does not exist"):
            validator.is_valid_file_type("nonexistent.pdf")

    def test_corrupt_image(self, tmp_path):
        """
        Test handling of corrupt image files.
        """
        # Create a corrupt PNG file
        corrupt_file = tmp_path / "corrupt.png"
        corrupt_file.write_bytes(b"\x89PNG\r\n\x1a\n" + b"corrupted data")

        validator = FileValidator()
        with pytest.raises(InvalidFileTypeError):
            validator.is_valid_file_type(corrupt_file)

    def test_extension_mismatch(self, tmp_path):
        """
        Test handling of files with mismatched extensions.
        """
        # Create a PNG image
        img = Image.new('RGB', (100, 100), color='red')
        # First save it as PNG
        png_path = tmp_path / "image.png"
        img.save(png_path)

        # Now create a copy with wrong extension
        wrong_ext_file = tmp_path / "same_image.pdf"
        with open(png_path, 'rb') as src:
            wrong_ext_file.write_bytes(src.read())

        validator = FileValidator()
        with pytest.raises(InvalidFileTypeError, match="extension does not match"):
            validator.is_valid_file_type(wrong_ext_file)

    def test_get_file_preview(self, tmp_path):
        """
        Test getting preview of files.
        """
        # Create a test image
        img = Image.new('RGB', (100, 100), color='red')
        img_path = tmp_path / "test.png"
        img.save(img_path)

        validator = FileValidator()
        preview = validator.get_file_preview(img_path)

        assert preview is not None
        assert isinstance(preview, dict)
        assert 'thumbnail' in preview  # For images
        assert 'mime_type' in preview
        assert preview['mime_type'] == 'image/png'

    def test_word_document_handling(self, tmp_path):
        """Test handling of Word documents."""
        try:
            from docx import Document

            # Create a test Word document
            doc = Document()
            doc.add_paragraph("Test content for Word document")
            doc.add_paragraph("Second paragraph")

            # Add a table
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Cell 1"
            table.cell(0, 1).text = "Cell 2"
            table.cell(1, 0).text = "Cell 3"
            table.cell(1, 1).text = "Cell 4"

            doc_path = tmp_path / "test.docx"
            doc.save(doc_path)

            # Test file validation
            validator = FileValidator()
            assert validator.is_valid_file_type(doc_path) == True
            assert validator.get_file_type(doc_path) == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

            # Test preview generation
            preview = validator.get_file_preview(doc_path)
            assert preview['preview_available'] == True
            assert "Test content" in preview['text_preview']

            # Test text extraction
            from src.classifier import ContentClassifier
            classifier = ContentClassifier()
            text = classifier.extract_text(doc_path)
            assert "Test content for Word document" in text
            assert "Second paragraph" in text
            assert "Cell 1" in text
            assert "Cell 4" in text
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_excel_document_handling(self, tmp_path):
        """Test handling of Excel documents."""
        try:
            from openpyxl import Workbook

            # Create a test Excel document
            wb = Workbook()
            ws = wb.active
            ws.title = "Test Sheet"
            ws['A1'] = "Header 1"
            ws['B1'] = "Header 2"
            ws['A2'] = "Data 1"
            ws['B2'] = "Data 2"

            excel_path = tmp_path / "test.xlsx"
            wb.save(excel_path)

            # Test file validation
            validator = FileValidator()
            assert validator.is_valid_file_type(excel_path) == True
            assert validator.get_file_type(excel_path) == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

            # Test preview generation
            preview = validator.get_file_preview(excel_path)
            assert preview['preview_available'] == True
            assert "Header 1" in preview['text_preview']
            assert "Data 1" in preview['text_preview']

            # Test text extraction
            from src.classifier import ContentClassifier
            classifier = ContentClassifier()
            text = classifier.extract_text(excel_path)
            assert "Header 1" in text
            assert "Header 2" in text
            assert "Data 1" in text
            assert "Data 2" in text
            assert "Test Sheet" in text
        except ImportError:
            pytest.skip("openpyxl not installed")

    def test_invalid_office_document(self, tmp_path):
        """Test handling of corrupted office documents."""
        # Create an invalid .docx file
        doc_path = tmp_path / "invalid.docx"
        doc_path.write_text("This is not a valid DOCX file")

        validator = FileValidator()
        with pytest.raises(InvalidFileTypeError):
            validator.is_valid_file_type(doc_path)

        # Create an invalid .xlsx file
        xlsx_path = tmp_path / "invalid.xlsx"
        xlsx_path.write_text("This is not a valid XLSX file")

        with pytest.raises(InvalidFileTypeError):
            validator.is_valid_file_type(xlsx_path)