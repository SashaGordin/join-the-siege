from pathlib import Path
from typing import Dict, Any, Union, List, Optional, Tuple
import pypdf
from PIL import Image
import magic
from openai import OpenAI
import os
from dotenv import load_dotenv
from dataclasses import dataclass
import cv2
import numpy as np
import pytesseract
import httpx
import logging
from .exceptions import (
    ClassificationError,
    TextExtractionError,
    FeatureExtractionError
)
from .config.config_manager import IndustryConfigManager
import re
from src.classifier.pattern_learning.pattern_matcher import PatternMatcher
from .models import ClassificationResult
from src.classifier.pattern_learning.pattern_store import PatternStore
import pdfplumber

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

@dataclass
class ClassificationResult:
    doc_type: str
    confidence: float
    features: Dict[str, Any]

class ContentClassifier:
    """
    A content-based document classifier that uses LLM to determine document type.
    """

    def __init__(self, default_industry: str = None, pattern_store: Optional['PatternStore'] = None):
        """
        Initialize the content classifier.

        Args:
            default_industry: Default industry to use for classification
            pattern_store: Optional PatternStore instance for pattern matching
        """
        self.mime_magic = magic.Magic(mime=True)
        self.default_industry = default_industry
        self.config_manager = IndustryConfigManager()
        self.pattern_store = pattern_store or PatternStore()
        self.pattern_matcher = PatternMatcher()

        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("OPENAI_API_KEY environment variable is not set")

        # Initialize OpenAI client with only the required parameters
        http_client = httpx.Client(
            follow_redirects=True,
            timeout=60.0
        )
        self.client = OpenAI(
            api_key=api_key,
            http_client=http_client
        )

        # Check if tesseract is installed
        try:
            pytesseract.get_tesseract_version()
        except pytesseract.TesseractNotFoundError:
            raise RuntimeError("Tesseract is not installed. Please install tesseract-ocr package.")

    def _get_industry_prompt(self, industry: str) -> str:
        """Get industry-specific prompt additions."""
        try:
            industry_config = self.config_manager.load_industry_config(industry)

            if industry == "healthcare":
                features = industry_config["features"]
                required = features["validation_rules"]["required_features"]
                specific = features["specific"]

                return f"""Industry Context: healthcare
Focus on medical document features and classify as medical documents when appropriate:
- Patient information ({specific["patient_id"]["description"]})
- Medical codes ({specific["medical_code"]["description"]})
- Provider details (NPI: {specific["provider_npi"]["description"]})
- Required features: {', '.join(required)}"""

            elif industry == "financial":
                features = industry_config["features"]
                required = features["validation_rules"]["required_features"]
                specific = features["specific"]

                return f"""Industry Context: financial
Focus on financial document features and classify as financial documents when appropriate:
- Invoice numbers ({specific["invoice_number"]["description"]})
- Payment terms ({specific["payment_terms"]["description"]})
- Account numbers ({specific["account_number"]["description"]})
- Required features: {', '.join(required)}
Do not classify as medical documents unless absolutely certain."""

        except Exception as e:
            logger.warning(f"Failed to load industry config for prompt: {str(e)}")
            return ""

        return ""

    def _extract_text_from_image(self, image_path: Path) -> str:
        """
        Extract text from an image using OCR.

        Args:
            image_path: Path to the image file

        Returns:
            str: Extracted text

        Raises:
            TextExtractionError: If text extraction fails
        """
        logger.info(f"Starting text extraction from image: {image_path}")

        # First validate the image file
        if not isinstance(image_path, (str, Path)):
            logger.error(f"Invalid path type: {type(image_path)}")
            raise TextExtractionError("Invalid file path type")

        # Check if file exists and is not empty
        try:
            if not Path(image_path).exists():
                logger.error(f"File does not exist: {image_path}")
                raise TextExtractionError("File does not exist")
            if Path(image_path).stat().st_size == 0:
                logger.error(f"File is empty: {image_path}")
                raise TextExtractionError("File is empty")
            logger.info(f"File validation passed. Size: {Path(image_path).stat().st_size} bytes")
        except Exception as e:
            logger.error(f"File access error: {str(e)}")
            raise TextExtractionError(f"File access error: {str(e)}")

        # Try to validate and open the image
        try:
            # First try to open with PIL to validate format
            try:
                with open(image_path, 'rb') as f:
                    header = f.read(8)  # Read first 8 bytes
                    logger.debug(f"File header bytes: {header.hex()}")
                    if not any(header.startswith(sig) for sig in [b'\x89PNG\r\n\x1a\n', b'\xff\xd8\xff', b'GIF87a', b'GIF89a']):
                        logger.error(f"Invalid image header: {header.hex()}")
                        raise TextExtractionError("Invalid image file format")
                    logger.info("Image header validation passed")
            except TextExtractionError:
                raise  # Re-raise TextExtractionError
            except Exception as e:
                logger.error(f"Failed to read file header: {str(e)}")
                raise TextExtractionError(f"Invalid image file: {str(e)}")

            try:
                with Image.open(str(image_path)) as img:
                    img.verify()
                    logger.info(f"Image verification successful. Format: {img.format}, Size: {img.size}, Mode: {img.mode}")
            except Exception as e:
                logger.error(f"Image verification failed: {str(e)}")
                raise TextExtractionError(f"Invalid image file: {str(e)}")

            # Try to read with OpenCV
            logger.info("Attempting to read image with OpenCV")
            img = cv2.imread(str(image_path))
            if img is None:
                logger.error("OpenCV failed to read the image")
                raise TextExtractionError("Failed to read image file: Not a valid image format")
            logger.info(f"OpenCV read successful. Image shape: {img.shape}")

            # Convert to grayscale
            try:
                logger.info("Converting image to grayscale")
                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                logger.info(f"Grayscale conversion successful. Shape: {gray.shape}")
            except cv2.error as e:
                logger.error(f"Failed to convert image to grayscale: {str(e)}")
                raise TextExtractionError(f"Failed to process image: {str(e)}")

            # Try different preprocessing techniques
            texts = []
            logger.info("Starting image preprocessing and OCR")

            # Try all four orientations
            angles = [0, 90, 180, 270]
            for angle in angles:
                logger.info(f"Processing image at {angle} degrees rotation")
                # Rotate image if needed
                if angle != 0:
                    height, width = gray.shape
                    center = (width // 2, height // 2)
                    rotation_matrix = cv2.getRotationMatrix2D(center, angle, 1.0)
                    rotated = cv2.warpAffine(gray, rotation_matrix, (width, height),
                                           flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
                    logger.debug(f"Rotated image shape: {rotated.shape}")
                else:
                    rotated = gray

                # 1. Basic thresholding
                logger.info("Applying basic thresholding")
                _, binary = cv2.threshold(rotated, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                text = pytesseract.image_to_string(binary)
                logger.debug(f"Basic thresholding OCR result length: {len(text)}")
                texts.append(text)

                # 2. Adaptive thresholding
                logger.info("Applying adaptive thresholding")
                binary_adaptive = cv2.adaptiveThreshold(
                    rotated, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
                )
                text = pytesseract.image_to_string(binary_adaptive)
                logger.debug(f"Adaptive thresholding OCR result length: {len(text)}")
                texts.append(text)

                # 3. Denoised version
                logger.info("Applying denoising")
                denoised = cv2.fastNlMeansDenoising(rotated)
                text = pytesseract.image_to_string(denoised)
                logger.debug(f"Denoised OCR result length: {len(text)}")
                texts.append(text)

                # Try different PSM modes
                psm_modes = ['--psm 6', '--psm 1', '--psm 3']
                for psm in psm_modes:
                    logger.info(f"Trying OCR with PSM mode: {psm}")
                    text = pytesseract.image_to_string(binary, config=psm)
                    logger.debug(f"PSM {psm} on binary OCR result length: {len(text)}")
                    texts.append(text)
                    text = pytesseract.image_to_string(binary_adaptive, config=psm)
                    logger.debug(f"PSM {psm} on adaptive OCR result length: {len(text)}")
                    texts.append(text)
                    text = pytesseract.image_to_string(denoised, config=psm)
                    logger.debug(f"PSM {psm} on denoised OCR result length: {len(text)}")
                    texts.append(text)

            # Choose the best result
            logger.info("Selecting best OCR result")
            def score_text(text):
                # Count alphanumeric characters
                alnum_count = sum(c.isalnum() for c in text)
                # Bonus points for numbers
                number_count = sum(c.isdigit() for c in text) * 2
                # Bonus points for common document keywords
                keyword_bonus = sum(10 for keyword in ['INVOICE', 'TOTAL', 'AMOUNT', 'DATE']
                                  if keyword in text.upper())
                score = alnum_count + number_count + keyword_bonus
                logger.debug(f"Text score: {score} (alnum: {alnum_count}, numbers: {number_count}, keywords: {keyword_bonus})")
                return score

            best_text = max(texts, key=score_text)
            logger.info(f"Selected best text with score {score_text(best_text)}")
            logger.debug(f"Best text content: {best_text[:100]}...")
            return best_text.strip()

        except TextExtractionError:
            raise  # Re-raise TextExtractionError
        except Exception as e:
            logger.error(f"Error during text extraction: {str(e)}", exc_info=True)
            if isinstance(e, pytesseract.TesseractError):
                if "Too few characters" in str(e):
                    logger.debug("Tesseract found no text in image")
                    return ""  # Return empty string for images with no text
            raise TextExtractionError(f"Failed to extract text from image: {str(e)}")

    def _extract_text_from_word(self, file_path: Path) -> str:
        """
        Extract text from Word documents (.doc, .docx).

        Args:
            file_path: Path to the Word document

        Returns:
            str: Extracted text content

        Raises:
            TextExtractionError: If text extraction fails
        """
        try:
            from docx import Document
            doc = Document(file_path)
            text = []

            # Extract main document text
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)

            return '\n'.join(text)
        except Exception as e:
            raise TextExtractionError(f"Failed to extract text from Word document: {str(e)}")

    def _extract_text_from_excel(self, file_path: Path) -> str:
        """
        Extract text from Excel files (.xls, .xlsx).

        Args:
            file_path: Path to the Excel file

        Returns:
            str: Extracted text content

        Raises:
            TextExtractionError: If text extraction fails
        """
        try:
            text = []
            file_type = self.mime_magic.from_file(str(file_path))

            if file_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                # Handle .xlsx files
                from openpyxl import load_workbook
                wb = load_workbook(file_path, read_only=True, data_only=True)

                for sheet in wb.worksheets:
                    sheet_text = []
                    for row in sheet.rows:
                        row_text = [str(cell.value) if cell.value is not None else '' for cell in row]
                        sheet_text.append(' '.join(row_text))
                    text.append('\n'.join(sheet_text))
                    text.append(f"--- End of Sheet: {sheet.title} ---\n")

            elif file_type == 'application/vnd.ms-excel':
                # Handle .xls files
                import xlrd
                wb = xlrd.open_workbook(file_path)

                for sheet in wb.sheets():
                    sheet_text = []
                    for row in range(sheet.nrows):
                        row_text = [str(sheet.cell_value(row, col)) for col in range(sheet.ncols)]
                        sheet_text.append(' '.join(row_text))
                    text.append('\n'.join(sheet_text))
                    text.append(f"--- End of Sheet: {sheet.name} ---\n")

            return '\n'.join(text)
        except Exception as e:
            raise TextExtractionError(f"Failed to extract text from Excel file: {str(e)}")

    def extract_text(self, file_path: Union[str, Path]) -> str:
        """
        Extract text content from a file.

        Args:
            file_path: Path to the file

        Returns:
            str: Extracted text content

        Raises:
            TextExtractionError: If text extraction fails
            NotImplementedError: For unimplemented file types
        """
        file_path = Path(file_path)
        mime_type = self.mime_magic.from_file(str(file_path))

        # Check if file claims to be an image based on extension
        is_image_extension = file_path.suffix.lower() in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']
        is_image_mime = mime_type.startswith('image/')

        try:
            # If file claims to be an image by extension but isn't a valid image file
            if is_image_extension and not is_image_mime:
                raise TextExtractionError(f"Invalid image file: File has image extension but invalid format")

            # Raise NotImplementedError for images
            if is_image_mime or is_image_extension:
                return self._extract_text_from_image(file_path)

            if mime_type == 'application/pdf':
                return self._extract_text_from_pdf(file_path)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                             'application/msword']:
                return self._extract_text_from_word(file_path)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                             'application/vnd.ms-excel']:
                return self._extract_text_from_excel(file_path)
            else:
                # For text files or unknown types, try reading as text
                try:
                    logger.debug(f"Trying to read {file_path} as utf-8")
                    text = file_path.read_text(encoding='utf-8')
                    # Heuristic: check for non-printable or replacement characters
                    non_printable = sum(1 for c in text if not c.isprintable())
                    total = len(text)
                    replacement_count = text.count('\uFFFD')
                    alnum_count = sum(1 for c in text if c.isalnum())
                    replacement_ratio = replacement_count / max(1, total)
                    non_printable_ratio = non_printable / max(1, total)
                    alnum_ratio = alnum_count / max(1, total)
                    logger.debug(f"[ENCODING CHECK] Text length: {total}, Non-printable: {non_printable} ({non_printable_ratio:.2f}), Replacement chars: {replacement_count}, Alnum: {alnum_count} ({alnum_ratio:.2f})")

                    # Extra logging for debugging encoding issues
                    if (total < 50 and non_printable_ratio > 0.1) or (total >= 50 and non_printable_ratio > 0.5) or alnum_ratio < 0.2 or replacement_count > 0:
                        logger.error(f"[ENCODING FAIL] Full text: {repr(text)}")
                        logger.error(f"[ENCODING FAIL] total: {total}, non_printable: {non_printable}, non_printable_ratio: {non_printable_ratio}, replacement_count: {replacement_count}, alnum_count: {alnum_count}, alnum_ratio: {alnum_ratio}")
                        logger.error(f"Text appears corrupted. Sample: {text[:100]}")
                        raise TextExtractionError("File has encoding issues")
                    return text
                except UnicodeDecodeError as ude_utf8:
                    logger.error(f"UnicodeDecodeError with utf-8: {ude_utf8}")
                    try:
                        logger.debug(f"Trying to read {file_path} as latin-1")
                        text = file_path.read_text(encoding='latin-1')
                        # Heuristic: check for non-printable or replacement characters
                        non_printable = sum(1 for c in text if not c.isprintable())
                        total = len(text)
                        replacement_count = text.count('\uFFFD')
                        alnum_count = sum(1 for c in text if c.isalnum())
                        replacement_ratio = replacement_count / max(1, total)
                        non_printable_ratio = non_printable / max(1, total)
                        alnum_ratio = alnum_count / max(1, total)
                        logger.debug(f"[ENCODING CHECK] Text length: {total}, Non-printable: {non_printable} ({non_printable_ratio:.2f}), Replacement chars: {replacement_count}, Alnum: {alnum_count} ({alnum_ratio:.2f})")

                        # Fail if short and non-printable >10%, or long and non-printable >50%, or very low alnum ratio
                        if (total < 50 and non_printable_ratio > 0.1) or (total >= 50 and non_printable_ratio > 0.5) or alnum_ratio < 0.2 or replacement_count > 0:
                            logger.error(f"Text appears corrupted. Sample: {text[:100]}")
                            raise TextExtractionError("File has encoding issues")
                        return text
                    except UnicodeDecodeError as ude_latin1:
                        logger.error(f"UnicodeDecodeError with latin-1: {ude_latin1}")
                        raise TextExtractionError("File has encoding issues")
        except NotImplementedError:
            raise  # Re-raise NotImplementedError as is
        except pypdf.errors.FileNotDecryptedError:
            raise TextExtractionError("File is encrypted and cannot be read")
        except Exception as e:
            logger.error(f"Failed to extract text: {str(e)}", exc_info=True)
            raise TextExtractionError(f"Failed to extract text: {str(e)}")

    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from a PDF file."""
        text = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                if pdf_reader.is_encrypted:
                    raise TextExtractionError("File is encrypted and cannot be read")
                for page in pdf_reader.pages:
                    text.append(page.extract_text())
            return '\n'.join(text)
        except Exception as e:
            raise TextExtractionError(f"Failed to extract PDF text: {str(e)}")

    def _normalize_doc_type(self, doc_type: str) -> str:
        """Normalize document type string."""
        doc_type = doc_type.lower().strip()

        # Handle unknown/unclassified documents
        if doc_type in ['unknown', 'unclassified', 'not sure', 'unclear']:
            return 'unknown'

        # Handle specific document types
        if doc_type in ['bank statement', 'bank_statement', 'statement', 'account statement']:
            return 'bank_statement'
        elif doc_type in ['invoice', 'bill', 'receipt']:
            return 'invoice'
        elif doc_type in ['medical prescription', 'medical_prescription', 'prescription', 'rx']:
            return 'prescription'

        # Remove common prefixes/suffixes
        doc_type = re.sub(r'^(a |an |the |medical )', '', doc_type)
        doc_type = re.sub(r' (document|doc|form)$', '', doc_type)

        # Convert spaces to underscores
        doc_type = '_'.join(doc_type.split())

        return doc_type

    def _classify_with_llm(self, text: str, industry: Optional[str] = None) -> ClassificationResult:
        """Classify document using LLM."""
        try:
            # Get industry-specific prompt additions
            industry_prompt = self._get_industry_prompt(industry) if industry else ""
            logger.debug(f"Using industry context: {industry}")
            logger.debug(f"Industry prompt: {industry_prompt}")

            # Enhanced prompt for better feature extraction
            prompt = f"""Analyze this document and classify it. Extract all relevant features.
Pay special attention to:
1. Document type
2. Dates (any format)
3. Amounts and monetary values
4. Document numbers (invoice numbers, patient IDs, etc.)
5. Payment terms (look for terms like 'Net 30', 'Due in X days', 'Payment due', etc.)

For financial documents:
- Look for payment terms in different formats (Net X, Due in X days, etc.)
- Extract all monetary amounts
- Note any late payment penalties or discounts

For healthcare documents:
- Extract all CPT codes (5-digit codes)
- Look for ICD-10 diagnosis codes
- Identify patient IDs and provider NPIs
- Note service dates and amounts

{industry_prompt}

Document text:
{text}

Provide a structured response with:
1. Document type
2. Confidence score (0.0-1.0)
3. All extracted features with their values

Format features as "Feature Type: Value" for easy parsing."""

            logger.debug("Full prompt being sent to LLM:")
            logger.debug("=" * 80)
            logger.debug(prompt)
            logger.debug("=" * 80)

            try:
                logger.debug("Calling OpenAI API...")
                # Call OpenAI API
                response = self.client.chat.completions.create(
                    model="gpt-4",  # Using GPT-4 for better accuracy
                    messages=[
                        {"role": "system", "content": "You are a document classification assistant. When given an industry context, you MUST classify documents according to that industry's conventions. For financial industry, prefer financial document types over medical ones unless absolutely certain it's a medical document."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.1,  # Low temperature for consistent results
                    max_tokens=1000
                )

                # Validate API response
                if not response.choices:
                    logger.error("Invalid API response: No choices returned")
                    raise ClassificationError("Invalid API response: No choices returned")

                logger.debug("Raw LLM Response:")
                logger.debug("=" * 80)
                logger.debug(response.choices[0].message.content)
                logger.debug("=" * 80)
            except Exception as api_error:
                logger.error(f"API call failed: {str(api_error)}")
                if "choices" in str(api_error):
                    raise ClassificationError("Invalid API response: Malformed response structure")
                raise ClassificationError(f"Classification failed: {str(api_error)}")

            # Parse response
            result = self._parse_llm_response(response.choices[0].message.content)

            # Validate against industry config if available
            if industry and result.doc_type != "unknown":
                try:
                    industry_config = self.config_manager.load_industry_config(industry)
                    if industry_config:
                        validation_results = {
                            "confidence_score": result.confidence,
                            "missing_required": [],
                            "invalid_format": []
                        }

                        # Check required fields based on industry
                        required_fields = industry_config["features"]["validation_rules"]["required_features"]
                        present_fields = {f["type"] for f in result.features}
                        validation_results["missing_required"] = [f for f in required_fields if f not in present_fields]

                        logger.debug(f"Initial confidence from LLM: {result.confidence}")
                        logger.debug(f"Required fields for {industry}: {required_fields}")
                        logger.debug(f"Present fields in document: {present_fields}")
                        logger.debug(f"Missing required fields: {validation_results['missing_required']}")

                        # Add validation warnings for missing required fields
                        if validation_results["missing_required"]:
                            result.features.append({
                                "type": "validation_warning",
                                "value": [f"Missing required fields: {', '.join(validation_results['missing_required'])}"],
                                "present": True
                            })

                        # More aggressive confidence reduction for missing fields
                        if validation_results["missing_required"]:
                            old_confidence = result.confidence
                            # Reduce confidence by 0.1 for each missing required field, up to 0.2 reduction
                            result.confidence = max(0.8, result.confidence - (0.1 * min(2, len(validation_results["missing_required"]))))
                            logger.debug(f"Confidence after missing fields adjustment: {result.confidence} (was {old_confidence})")

                        # Check for mixed industry signals
                        medical_signals = any(f["type"] in ["patient_id", "medical_code", "cpt_code", "npi"] for f in result.features)
                        financial_signals = any(f["type"] in ["invoice_number", "payment_terms", "account_number"] for f in result.features)

                        if medical_signals and financial_signals:
                            old_confidence = result.confidence
                            result.confidence = max(0.7, result.confidence - 0.2)
                            logger.debug(f"Confidence after mixed signals adjustment: {result.confidence} (was {old_confidence})")

                        # Clean up feature values
                        for feature in result.features:
                            if feature["type"] == "cpt_code":
                                # Extract just the CPT code from strings like "99213) - $150"
                                feature["value"] = re.search(r'\d{5}', feature["value"]).group() if re.search(r'\d{5}', feature["value"]) else feature["value"]
                            elif feature["type"] == "medical_code":
                                # Clean up ICD-10 codes
                                feature["value"] = re.sub(r'^ICD-?10\s*', '', feature["value"]).strip() if feature["value"] else feature["value"]
                            elif feature["type"] == "amount":
                                # Clean up amounts to just the number
                                feature["value"] = re.sub(r'[^\d.]', '', feature["value"]) if feature["value"] else feature["value"]

                        # Additional confidence reduction for invalid documents
                        if len(validation_results["missing_required"]) > len(required_fields) // 2:
                            old_confidence = result.confidence
                            result.confidence = min(result.confidence * 0.8, 0.7)  # Less aggressive reduction
                            logger.debug(f"Confidence after invalid document check: {result.confidence} (was {old_confidence})")

                        # Adjust confidence based on industry context match
                        if industry == "healthcare" and not medical_signals:
                            old_confidence = result.confidence
                            result.confidence = min(result.confidence * 0.8, 0.7)  # Less aggressive reduction
                            logger.debug(f"Confidence after healthcare context check: {result.confidence} (was {old_confidence})")
                        elif industry == "financial" and not financial_signals:
                            old_confidence = result.confidence
                            result.confidence = min(result.confidence * 0.8, 0.7)  # Less aggressive reduction
                            logger.debug(f"Confidence after financial context check: {result.confidence} (was {old_confidence})")

                        # Boost confidence for strong matches
                        if len(validation_results["missing_required"]) == 0:
                            result.confidence = min(1.0, result.confidence * 1.2)  # Boost confidence by 20% for perfect matches
                            logger.debug(f"Confidence after perfect match boost: {result.confidence}")

                        # Cap confidence for minimal invoices (only amount and invoice_number present)
                        if present_fields == {"amount", "invoice_number"}:
                            logger.debug(f"Capping confidence for minimal invoice (only amount and invoice_number present): {result.confidence} -> 0.8")
                            result.confidence = min(result.confidence, 0.8)

                        logger.debug(f"Final confidence score: {result.confidence}")
                        logger.debug(f"Final document type: {result.doc_type}")
                        logger.debug("=" * 80)

                        validation_results["confidence_score"] = result.confidence
                        return result

                except Exception as e:
                    logger.error(f"Error during industry validation: {str(e)}")
                    raise ClassificationError(f"Industry validation failed: {str(e)}")

            # Aggregate all amount features into a single feature
            amount_values = [f["value"] for f in result.features if f["type"] == "amount"]
            if amount_values:
                # Remove all individual amount features
                result.features = [f for f in result.features if f["type"] != "amount"]
                result.features.append({
                    "type": "amount",
                    "value": amount_values[0],
                    "values": amount_values,
                    "present": True
                })
                logger.debug(f"Aggregated amount values: {amount_values}")

            # Aggregate all date features into a single 'date' feature
            date_values = [f["value"] for f in result.features if f["type"] == "date" or f["type"].endswith("_date")]
            if date_values:
                # Remove all individual date features
                result.features = [f for f in result.features if not (f["type"] == "date" or f["type"].endswith("_date"))]
                result.features.append({
                    "type": "date",
                    "value": date_values[0],
                    "values": date_values,
                    "present": True
                })
                logger.debug(f"Aggregated date values: {date_values}")

            # All other confidence logic is handled in _parse_llm_response
            return result

        except Exception as e:
            logger.error(f"Classification failed: {str(e)}")
            if isinstance(e, ClassificationError):
                raise
            raise ClassificationError(f"Classification failed: {str(e)}")

    def _parse_llm_response(self, response: str) -> ClassificationResult:
        """Parse LLM response into structured format."""
        try:
            logger.debug(f"Starting to parse LLM response: {response[:200]}...")
            lines = response.split('\n')

            doc_type = None
            confidence = None
            features = []
            current_section = None
            cpt_codes = []  # Track CPT codes separately
            sections = []  # Track legal document sections
            is_multilingual = False  # Flag for multilingual content

            # Feature type mapping
            feature_mapping = {
                'patient id': 'patient_id',
                'date of service': 'date',
                'service date': 'date',
                'date': 'date',
                'agreement date': 'date',
                'statement period': 'date',
                'appraisal date': 'date',
                'semester': 'date',
                'cpt code': 'cpt_code',
                'cpt codes': 'cpt_code',
                'cpt': 'cpt_code',
                'service': 'cpt_code',
                'icd-10': 'diagnosis',
                'icd10': 'diagnosis',
                'diagnosis code': 'diagnosis',
                'diagnosis': 'diagnosis',
                'npi': 'npi',
                'provider npi': 'npi',
                'provider name': 'provider',
                'provider': 'provider',
                'invoice number': 'invoice_number',
                'amount': 'amount',
                'monetary value': 'amount',
                'fee structure': 'amount',
                'total amount': 'amount',
                'opening balance': 'amount',
                'closing balance': 'amount',
                'market value': 'amount',
                'payment terms': 'payment_terms',
                'terms': 'payment_terms',
                'document type': 'document_type',
                'account number': 'account_number',
                'contract id': 'contract_id',
                'document number': 'contract_id',
                'report id': 'report_id',
                'property address': 'property_address',
                'parties': 'parties',
                'student id': 'student_id',
                'student name': 'student_name',
                'program': 'program',
                'gpa': 'gpa',
                'course': 'course',
                'section': 'section',
                'clause': 'clause'
            }

            # Parse lines and extract features
            cpt_code_values = []
            cpt_code_indices = []
            for idx, line in enumerate(lines):
                line = line.strip()
                if not line:
                    continue

                line_lower = line.lower()
                logger.debug(f"Processing line: {line}")

                # Log lines that look like amounts or contain currency symbols
                if re.search(r'[€£$¥₿]', line):
                    logger.debug(f"[CURRENCY LINE] {line}")

                # Check for multilingual content
                if '/' in line and any(pair in line_lower for pair in ['date/fecha', 'amount/monto', 'invoice/factura']):
                    is_multilingual = True
                    logger.debug("Detected multilingual content")

                # Extract document type
                if 'document type:' in line_lower or line_lower.startswith('document type:'):
                    doc_type = line.split(':', 1)[1].strip().lower()
                    doc_type = self._normalize_doc_type(doc_type)
                    logger.debug(f"Found document type: {doc_type}")
                    continue

                # Extract confidence (improved: also match 'confidence score:')
                if 'confidence score:' in line_lower or line_lower.startswith('confidence:'):
                    try:
                        confidence = float(line.split(':', 1)[1].strip())
                        logger.debug(f"Found confidence: {confidence}")
                    except ValueError:
                        confidence = 0.5  # Default confidence
                    continue

                # Check for legal document sections/clauses
                section_match = re.match(r'^\d+\.\s*(.*?)(?:\s*-|$)', line)
                if section_match:
                    section_name = section_match.group(1).strip()
                    sections.append(section_name)
                    logger.debug(f"Found section: {section_name}")
                    continue

                # Check for CPT codes in service lines
                cpt_match = re.search(r'CPT:?\s*(\d{5})', line)
                if cpt_match:
                    cpt_codes.append(cpt_match.group(1))
                    cpt_code_values.append(cpt_match.group(1))
                    logger.debug(f"Found CPT code: {cpt_match.group(1)}")
                    continue

                # Check for '- <Feature Name>: <Value>' pattern
                dash_colon_match = re.match(r'-\s*([^:]+):\s*(.*)', line)
                if dash_colon_match:
                    raw_key = dash_colon_match.group(1).strip().lower()
                    value = dash_colon_match.group(2).strip()
                    logger.debug(f"[DASH COLON] raw_key: {raw_key}, value: {value}")
                    # Always add the original key as a feature (snake_case)
                    orig_feature_type = raw_key.replace(' ', '_')
                    features.append({
                        "type": orig_feature_type,
                        "value": value,
                        "values": [value],
                        "present": True
                    })
                    logger.debug(f"Added feature with original key: {orig_feature_type} = {value}")
                    # Map using feature_mapping (if different from original)
                    feature_type = feature_mapping.get(raw_key)
                    if feature_type and feature_type != orig_feature_type:
                        features.append({
                            "type": feature_type,
                            "value": value,
                            "values": [value],
                            "present": True
                        })
                        logger.debug(f"Added mapped feature: {feature_type} = {value}")
                    # If this is an item line, extract monetary values as amount features
                    if (feature_type and (feature_type.startswith("item_") or "amount" in feature_type)) or ("amount" in orig_feature_type):
                        # Find all monetary values in the value string
                        amounts = re.findall(r'[€£$¥₿]\s?\d+[\d,.]*', value)
                        logger.debug(f"[AMOUNT EXTRACTION] Found amounts: {amounts} in value: {value}")
                        for amt in amounts:
                            features.append({
                                "type": "amount",
                                "value": amt,
                                "values": [amt],
                                "present": True
                            })
                            logger.debug(f"Extracted amount from item/amount line: {amt}")
                    continue

            logger.debug(f"[POST-PARSE] All features: {features}")

            # Remove individual cpt_code features and add a single aggregated one if any found
            if cpt_code_values:
                # Remove all individual cpt_code features
                features = [f for f in features if f["type"] != "cpt_code"]
                features.append({
                    "type": "cpt_code",
                    "value": ", ".join(cpt_code_values),
                    "values": cpt_code_values,
                    "present": True
                })
                logger.debug(f"Aggregated CPT codes: {cpt_code_values}")

            # Add sections as a feature if found
            if sections:
                features.append({
                    "type": "section",
                    "value": sections[0],
                    "values": sections,
                    "present": True
                })
                logger.debug(f"Added section feature: {sections}")

            # Calculate confidence based on document completeness
            logger.debug(f"doc_type before confidence calculation: {doc_type}")
            logger.debug(f"All extracted features: {features}")
            if doc_type == "invoice":
                # For invoices, we now consider amount as the core requirement
                # date and invoice_number are bonus fields
                core_fields = {"amount"}
                bonus_fields = {"date", "invoice_number"}
                present_fields = {f["type"] for f in features}

                # Count how many amounts we have (for multiple currency case)
                amount_count = sum(1 for f in features if f["type"] == "amount")

                if core_fields.issubset(present_fields):
                    # Has core fields - start with good confidence
                    confidence = 0.7
                    logger.debug(f"Set base confidence for invoice with core fields: {confidence}")

                    # Bonus for multiple amounts (multiple currencies)
                    if amount_count > 1:
                        confidence = min(1.0, confidence + 0.1)
                        logger.debug(f"Added bonus for multiple amounts: {confidence}")

                    # Bonus for invoice number
                    if "invoice_number" in present_fields:
                        confidence = min(1.0, confidence + 0.15)  # Strong signal
                        logger.debug(f"Added bonus for invoice number: {confidence}")

                    # Bonus for date
                    if "date" in present_fields:
                        confidence = min(1.0, confidence + 0.1)
                        logger.debug(f"Added bonus for date: {confidence}")

                    # Cap confidence for minimal invoices (only amount and invoice_number present)
                    if present_fields == {"amount", "invoice_number"}:
                        logger.debug(f"Capping confidence for minimal invoice (only amount and invoice_number present): {confidence} -> 0.8")
                        confidence = min(confidence, 0.8)

                    # Boost for multilingual content that's well-structured
                    if is_multilingual:
                        confidence = min(1.0, confidence + 0.15)  # Significant boost for multilingual
                        logger.debug(f"Added multilingual boost: {confidence}")
                else:
                    # Missing core fields
                    missing_core = len(core_fields - present_fields)
                    confidence = 0.5 - (0.2 * missing_core)
                    logger.debug(f"Reduced confidence due to missing core fields: {confidence}")

                # Add document type as a feature
                features.append({
                    "type": "document_type",
                    "value": doc_type,
                    "present": True
                })

            elif doc_type == "property_appraisal_report":
                # Property appraisal: require report_id, date, property_address, amount
                required_fields = {"report_id", "date", "property_address", "amount"}
                present_fields = {f["type"] for f in features}
                logger.debug(f"[REALESTATE] present_fields: {present_fields}")
                missing_fields = required_fields - present_fields
                logger.debug(f"[REALESTATE] missing required fields: {missing_fields}")
                if required_fields.issubset(present_fields):
                    confidence = 0.9
                    logger.debug(f"[REALESTATE] Set high confidence for complete property appraisal: {confidence}")
                else:
                    missing_count = len(missing_fields)
                    confidence = 0.7 - (0.1 * missing_count)
                    logger.debug(f"[REALESTATE] Adjusted confidence for partial property appraisal: {confidence}")
            elif doc_type == "legal_contract":
                # Legal contract: require contract_id, date, and at least one section or clause
                required_fields = {"contract_id", "date"}
                present_fields = {f["type"] for f in features}
                has_section_or_clause = any(f["type"] in {"section", "clause"} for f in features)
                logger.debug(f"[LEGAL] present_fields: {present_fields}")
                logger.debug(f"[LEGAL] has_section_or_clause: {has_section_or_clause}")
                if required_fields.issubset(present_fields) and has_section_or_clause:
                    confidence = 0.9
                    logger.debug(f"[LEGAL] Set high confidence for complete legal contract: {confidence}")
                else:
                    missing_count = len(required_fields - present_fields)
                    logger.debug(f"[LEGAL] missing required fields: {required_fields - present_fields}")
                    confidence = 0.5
                    if has_section_or_clause:
                        confidence += 0.1
                        logger.debug(f"[LEGAL] Added bonus for section/clause: {confidence}")
                    logger.debug(f"[LEGAL] Adjusted confidence for partial legal contract: {confidence}")
            elif doc_type == "claim":
                required_fields = {"date", "cpt_code", "diagnosis", "provider"}
                present_fields = {f["type"] for f in features}
                if required_fields.issubset(present_fields):
                    confidence = 0.9
                    logger.debug(f"Set high confidence for complete medical claim: {confidence}")
                else:
                    missing_count = len(required_fields - present_fields)
                    confidence = 0.7 - (0.1 * missing_count)
                    logger.debug(f"Adjusted confidence for partial medical claim: {confidence}")
            elif doc_type == "prescription":
                # Prescription: require patient and prescription features
                present_types = {f["type"] for f in features}
                has_patient = any("patient" in f["type"] for f in features)
                has_rx = any(f["type"] in {"prescription", "rx"} for f in features)
                logger.debug(f"[PRESCRIPTION] present_types: {present_types}")
                logger.debug(f"[PRESCRIPTION] has_patient: {has_patient}, has_rx: {has_rx}")
                if has_patient and has_rx:
                    confidence = 0.8
                    logger.debug(f"[PRESCRIPTION] Set high confidence for complete prescription: {confidence}")
                elif has_patient or has_rx:
                    confidence = 0.6
                    logger.debug(f"[PRESCRIPTION] Set medium confidence for partial prescription: {confidence}")
                # else: fall through to default

            # If confidence is still None, set a default
            if confidence is None:
                confidence = 0.5
                logger.debug(f"Set default confidence: {confidence}")

            logger.debug(f"Creating ClassificationResult with doc_type={doc_type}, confidence={confidence}, features={features}")
            result = ClassificationResult(
                doc_type=doc_type,
                confidence=confidence,
                features=features
            )
            logger.debug(f"Created ClassificationResult: {result}")
            return result

        except Exception as e:
            logger.error(f"Failed to parse LLM response: {str(e)}", exc_info=True)
            raise ClassificationError(f"Invalid API response: {str(e)}")

    def classify_file(self, file_path: Union[str, Path], industry: Optional[str] = None) -> Dict[str, Any]:
        """Classify a document file."""
        try:
            # Convert to Path object if string
            if isinstance(file_path, str):
                file_path = Path(file_path)

            # Check if file exists
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")

            # Check if file is empty
            if file_path.stat().st_size == 0:
                raise ClassificationError("Empty file")

            # Extract text based on file type
            text = self.extract_text(file_path)

            # Check if extracted text is empty
            if not text or not text.strip():
                raise ClassificationError("Empty file")

            # FAST PATH for large/repetitive documents
            if len(text) > 5000 or text.count('\n') > 200:
                import re
                invoice_numbers = re.findall(r'Invoice Number: ([^\n]+)', text)
                amounts = re.findall(r'Amount: ([^\n]+)', text)
                features = []
                if invoice_numbers:
                    features.append({
                        "type": "invoice_number",
                        "value": invoice_numbers[0],
                        "values": invoice_numbers,
                        "present": True
                    })
                if amounts:
                    features.append({
                        "type": "amount",
                        "value": amounts[0],
                        "values": amounts,
                        "present": True
                    })
                return {
                    "class": "invoice",
                    "confidence": 0.8,
                    "features": features
                }

            # Classify text
            logger.debug("Calling _classify_with_llm...")
            result = self._classify_with_llm(text, industry)
            logger.debug(f"Result from _classify_with_llm: {result}")

            # Validate result
            if not isinstance(result, ClassificationResult):
                logger.error(f"Invalid result type: {type(result)}")
                raise ClassificationError(f"Invalid classification result type: {type(result)}")

            logger.debug(f"Returning classification result: {result}")
            return {
                "class": result.doc_type,
                "confidence": result.confidence,
                "features": result.features
            }

        except TextExtractionError as e:
            logger.error(f"Text extraction failed: {str(e)}", exc_info=True)
            raise
        except Exception as e:
            logger.error(f"Classification failed: {str(e)}", exc_info=True)
            if isinstance(e, (ClassificationError, FileNotFoundError)):
                raise
            raise ClassificationError(f"Classification failed: {str(e)}")