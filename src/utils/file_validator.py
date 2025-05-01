import magic
from pathlib import Path
from typing import Union, Dict, Any
import io
from PIL import Image

from .exceptions import (
    InvalidFileTypeError,
    FileTooLargeError,
    FileCorruptError,
    FileAccessError,
)

class FileValidator:
    """
    Validates files for type, size, and content.
    Provides preview functionality for supported file types.
    """

    # Maximum file size (10MB)
    MAX_FILE_SIZE = 10 * 1024 * 1024

    # Maximum preview size for images
    MAX_PREVIEW_SIZE = (300, 300)

    # Allowed MIME types and their corresponding extensions
    ALLOWED_MIME_TYPES = {
        'application/pdf': '.pdf',
        'image/jpeg': '.jpg',
        'image/png': '.png',
        'text/plain': '.txt',
        # Common variations of text MIME types
        'text/x-python': '.txt',  # Python files are text
        'text/x-java': '.txt',    # Java files are text
        'text/x-c': '.txt',       # C files are text
        'text/x-cpp': '.txt',     # C++ files are text
        'text/x-script': '.txt',  # Script files are text
        # Office document formats
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': '.xlsx',
        'application/msword': '.doc',
        'application/vnd.ms-excel': '.xls',
    }

    def __init__(self):
        """Initialize the file validator."""
        self.magic = magic.Magic(mime=True)

    def get_file_type(self, file_path: Union[str, Path]) -> str:
        """
        Get the MIME type of a file.

        Args:
            file_path: Path to the file

        Returns:
            str: MIME type of the file

        Raises:
            FileAccessError: If file cannot be accessed
        """
        file_path = Path(file_path)
        try:
            return self.magic.from_file(str(file_path))
        except Exception as e:
            raise FileAccessError(f"Cannot access file: {e}")

    def is_valid_file_type(self, file_path: Union[str, Path]) -> bool:
        """
        Check if the file is of a valid type and size.

        Args:
            file_path: Path to the file

        Returns:
            bool: True if file is valid

        Raises:
            InvalidFileTypeError: If file type is not supported
            FileTooLargeError: If file is too large
            FileAccessError: If file cannot be accessed
        """
        file_path = Path(file_path)

        # Check if file exists
        if not file_path.exists():
            raise FileAccessError("File does not exist")

        # Check if file is empty
        if file_path.stat().st_size == 0:
            raise InvalidFileTypeError("Empty file")

        # Check file size
        if file_path.stat().st_size > self.MAX_FILE_SIZE:
            raise FileTooLargeError("File exceeds maximum allowed size")

        # Get actual MIME type from file content
        actual_mime_type = self.get_file_type(file_path)
        actual_ext = file_path.suffix.lower()

        # For text files, normalize the MIME type and check extension
        if actual_mime_type.startswith('text/'):
            if actual_ext in ['.txt', '.text', '.log']:
                return True
            # If it's a text file with wrong extension, treat it as unsupported
            raise InvalidFileTypeError(f"Unsupported file type: text file with invalid extension {actual_ext}")

        # For images, allow them through validation (they'll be handled in classification)
        if actual_mime_type in ['image/jpeg', 'image/png']:
            if actual_ext in ['.jpg', '.jpeg', '.png']:
                return True
            raise InvalidFileTypeError(f"File extension does not match content type")

        # For PDFs, check both MIME type and extension
        if actual_mime_type == 'application/pdf':
            if actual_ext == '.pdf':
                return True
            raise InvalidFileTypeError(f"File extension does not match content type")

        # For Word documents
        if actual_mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            if actual_ext == '.docx':
                return True
            raise InvalidFileTypeError(f"File extension does not match content type")

        # For Excel documents
        if actual_mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
            if actual_ext == '.xlsx':
                return True
            raise InvalidFileTypeError(f"File extension does not match content type")

        # If we get here, the file type is not supported
        raise InvalidFileTypeError(f"Unsupported file type: {actual_mime_type}")

    def get_file_preview(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Get a preview of the file contents.

        Args:
            file_path: Path to the file

        Returns:
            Dict containing preview information
        """
        file_path = Path(file_path)
        mime_type = self.get_file_type(file_path)

        preview = {
            'preview_available': False,
            'mime_type': mime_type,
            'text_preview': None,
            'metadata': {}
        }

        try:
            # Handle different file types
            if mime_type.startswith('text/'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    preview['text_preview'] = f.read(1000)  # First 1000 chars
                    preview['preview_available'] = True

            elif mime_type == 'application/pdf':
                # Use PyPDF2 for PDF preview
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                if len(reader.pages) > 0:
                    preview['text_preview'] = reader.pages[0].extract_text()[:1000]
                    preview['preview_available'] = True
                    preview['metadata']['pages'] = len(reader.pages)

            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Use python-docx for Word documents
                from docx import Document
                doc = Document(file_path)
                text = []
                for para in doc.paragraphs[:5]:  # First 5 paragraphs
                    text.append(para.text)
                preview['text_preview'] = '\n'.join(text)[:1000]
                preview['preview_available'] = True
                preview['metadata']['paragraphs'] = len(doc.paragraphs)

            elif mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                # Use openpyxl for Excel documents
                from openpyxl import load_workbook
                wb = load_workbook(file_path, read_only=True)
                text = []
                sheet = wb.active
                for row in list(sheet.rows)[:5]:  # First 5 rows
                    text.append(' | '.join(str(cell.value) for cell in row if cell.value))
                preview['text_preview'] = '\n'.join(text)[:1000]
                preview['preview_available'] = True
                preview['metadata']['sheets'] = len(wb.sheetnames)
                wb.close()

            elif mime_type in ['image/jpeg', 'image/png']:
                # Generate both thumbnail and OCR text
                img = Image.open(file_path)

                # Create thumbnail
                img.thumbnail(self.MAX_PREVIEW_SIZE)
                thumb_io = io.BytesIO()
                img.save(thumb_io, format=img.format)
                preview['thumbnail'] = thumb_io.getvalue()

                # Add image metadata
                preview['metadata'].update({
                    'width': img.width,
                    'height': img.height,
                    'format': img.format
                })

                # Extract text using OCR
                import pytesseract
                preview['text_preview'] = pytesseract.image_to_string(img)[:1000]
                preview['preview_available'] = True

        except Exception as e:
            logger.warning(f"Failed to generate preview for {file_path}: {str(e)}")
            preview['error'] = str(e)

        return preview